"""FastAPI-мост к LangGraph-пайплайну для фронтенда.

Возможности:
- старт прогона (тема/жанр/язык/порция глав/перевод вкл-выкл);
- live-прогресс агентов (поллинг state + SSE);
- ручная правка артефактов на любом этапе (update_state);
- per-stage «попроси ИИ переделать» (revise) + правка плана конкретной главы;
- структура порциями: «ещё порцию» / «перейти к написанию»;
- решения по findings редактора: принять/отклонить/комментировать + авто-judge;
- откат на предыдущий этап (rollback);
- step-mode: пауза после каждого бота + resume.

Запуск: uvicorn server.app:app --reload --port 8000
"""
from __future__ import annotations

import asyncio
import json
import os
import queue
import threading
import uuid
from dataclasses import dataclass, field
from typing import Any, Optional

from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from narrative import nodes
from narrative.graph import STAGE_NODES, build_graph, sqlite_saver
from narrative.llm import LLMLimitError
from narrative.config import CHAPTER_WORDS_DEFAULT
from narrative.routing import apply_decisions
from narrative.state import Chapter, JudgeOut

load_dotenv()
app = FastAPI(title="ai-narrative API")
# Узкий CORS: фронт ходит через Next-прокси (same-origin), напрямую в браузере
# CORS не нужен. Разрешаем только локальные dev-origin'ы — это режет drive-by
# CSRF с произвольных сайтов на мутирующие/токен-эндпоинты (бэк на 127.0.0.1).
_ALLOWED = [o for o in os.environ.get(
    "CORS_ORIGINS",
    "http://localhost:3030,http://127.0.0.1:3030,"
    "http://localhost:3000,http://127.0.0.1:3000",
).split(",") if o]
app.add_middleware(
    CORSMiddleware, allow_origins=_ALLOWED,
    allow_methods=["*"], allow_headers=["*"],
)

_SAVER = sqlite_saver()
_GRAPH_AUTO = build_graph(_SAVER)
_GRAPH_STEP = build_graph(_SAVER, interrupt_after=STAGE_NODES)
_LOCK = threading.Lock()

# Прямой вызов узла (для per-stage revise / re-check) — узлы чистые (state→dict).
NODE_FUNCS = {
    "logline": nodes.logline_node,
    "synopsis": nodes.synopsis_node,
    "characters": nodes.characters_node,
    "locations": nodes.locations_node,
    "structure": nodes.structure_node,
    "dialogue": nodes.dialogue_node,
    "adult": nodes.adult_node,
    "editor": nodes.editor_node,
    "translation": nodes.translation_node,
}

# Откат: какой узел «проиграть как предыдущий», чтобы next встал на нужный этап,
# и какие (не-reducer) поля очистить.
ROLLBACK = {
    "synopsis":   ("logline",    ["synopsis", "characters", "locations", "chapters", "chapter_idx", "structure_done"]),
    "characters": ("synopsis",   ["characters", "locations", "chapters", "chapter_idx", "structure_done"]),
    "locations":  ("characters", ["locations", "chapters", "chapter_idx", "structure_done"]),
    "structure":  ("characters", ["chapters", "chapter_idx", "structure_done"]),
    "dialogue":   ("structure",  ["chapter_idx"]),  # перезапуск написания глав с 0
}


@dataclass
class Run:
    thread_id: str
    step_mode: bool
    events: queue.Queue = field(default_factory=queue.Queue)
    status: str = "idle"
    printed: int = 0
    worker: Optional[threading.Thread] = None
    error: str = ""  # текст последней ошибки воркера (для UI)
    limit: Optional[dict] = None  # инфо об исчерпании лимита (#5): баннер выбора
    gen: Optional[dict] = None  # live-партиал текущей генерации {stage, idx, text}
    cancel: bool = False  # запрос остановки текущей генерации (любой этап)
    structure_dirty: bool = False  # главы менялись вручную → предложить проверку


class _Cancelled(BaseException):
    """Сигнал «Стоп»: наследуем BaseException, чтобы НЕ быть проглоченным
    `except Exception` (в live.feed-синке и ретраях LLM) — иначе отмена терялась
    и кнопка «Стоп» не срабатывала на стрим-этапах."""


RUNS: dict[str, Run] = {}


def _graph(run: Run):
    return _GRAPH_STEP if run.step_mode else _GRAPH_AUTO


def _config(thread_id: str) -> dict:
    return {"configurable": {"thread_id": thread_id}}


def _get_run(thread_id: str) -> Run:
    run = RUNS.get(thread_id)
    if run is not None:
        return run
    # Прогона нет в памяти (рестарт бэка / краш процесса), но state мог
    # сохраниться в SQLite-чекпоинте → восстанавливаем Run, чтобы продолжить
    # с места обрыва, а не терять весь прогресс.
    try:
        snap = _GRAPH_STEP.get_state(_config(thread_id))
    except Exception:
        snap = None
    if snap and (snap.values or snap.next):
        status = "done" if not snap.next else "paused"
        run = Run(thread_id=thread_id, step_mode=True, status=status)
        RUNS[thread_id] = run
        return run
    raise HTTPException(404, "run not found")


def _bg_op(run: Run, fn, stage: str | None = None,
           idx: Optional[int] = None) -> dict:
    """Запускает ручную LLM-операцию В ФОНЕ (LLM долгий — иначе Next-прокси рвёт
    по таймауту → 500). Возвращает сразу; фронт поллит status.

    stage задаёт, КАКОЙ этап перегенерится (для корректного индикатора и live —
    иначе UI показал бы next-узел графа). Привязывает live-sink в этом потоке,
    чтобы ревизия стримилась как обычная генерация.

    Не ломает позицию графа: по завершении/ошибке возвращаем прежний статус
    (paused/done), ошибку кладём в run.error, прогресс в чекпоинте цел."""
    from narrative import live  # noqa: PLC0415
    prev = run.status if run.status in ("paused", "done") else "paused"

    def _worker():
        if stage:
            run.gen = {"stage": stage, "idx": idx, "text": ""}
            def _sink(_stage, _idx, text):  # стрим ревизии во фронт
                if run.cancel:
                    raise _Cancelled()
                run.gen = {"stage": stage, "idx": idx, "text": text}
                run.events.put({"type": "gen", "stage": stage, "idx": idx, "text": text})
            live.bind(_sink, stage, idx)
        run.events.put({"type": "status", "status": "running"})
        try:
            fn()
            run.status = prev
            run.events.put({"type": "status", "status": prev})
        except _Cancelled:  # «Стоп» во время ручной операции
            run.cancel = False
            run.status = prev
            run.events.put({"type": "status", "status": prev})
        except LLMLimitError as exc:  # #5: лимит → баннер выбора, прогресс цел
            run.limit = {"provider": exc.provider, "reset_at": exc.reset_at,
                         "kind": exc.kind, "message": str(exc)[:200]}
            run.status = "limit"
            run.events.put({"type": "status", "status": "limit",
                            "limit": run.limit})
        except Exception as exc:  # noqa: BLE001
            run.error = f"Операция не удалась (LLM/провайдер): {str(exc)[:200]}"
            run.status = prev
            run.events.put({"type": "status", "status": prev,
                            "error": run.error})
        finally:
            live.clear()
            run.gen = None

    # старт под _LOCK: повторная проверка занятости закрывает TOCTOU-гонку
    # (двойной клик → два одновременных LLM-воркера → потерянные записи).
    # Статус running ставим СИНХРОННО здесь же (после проверки!) — иначе
    # немедленный poll фронта ловит «paused» и перестаёт поллить (правки
    # были видны только после F5).
    with _LOCK:
        if run.worker and run.worker.is_alive():
            raise HTTPException(409, "run is busy")
        run.cancel = False
        run.status = "running"
        run.error = ""
        run.worker = threading.Thread(target=_worker, daemon=True)
        run.worker.start()
    return {"ok": True, "started": True}


def _state(thread_id: str) -> dict:
    run = _get_run(thread_id)
    return dict(_graph(run).get_state(_config(thread_id)).values or {})


# кэш перевода замечаний редактора EN→RU (для показа; правки в коде остаются EN)
_findings_ru_cache: dict[str, str] = {}
# circuit breaker: Google недоступен (datacenter-IP блок и т.п.) → не дёргаем его
# на каждый poll /state, иначе первый запрос после рестарта висит минутами.
_gt_down_until: float = 0.0


def _to_ru_display(text: str) -> str:
    """Перевести текст замечания на русский ДЛЯ ПОКАЗА (быстрый Google, кэш).
    Внутренняя логика (ревизии/судья) использует оригинальный английский.

    /state поллится каждые ~1с — здесь НЕЛЬЗЯ висеть: таймаут короткий (3с),
    при сбое Google отключается на 5 минут (показываем английский как есть)."""
    global _gt_down_until  # noqa: PLW0603
    import time as _time  # noqa: PLC0415
    t = (text or "").strip()
    if not t:
        return text
    if t in _findings_ru_cache:
        return _findings_ru_cache[t]
    if _time.time() < _gt_down_until:
        return text  # Google лежит — не кэшируем, попробуем после паузы
    try:
        from narrative.gtrans import google_translate  # noqa: PLC0415
        ru = google_translate(t, "ru", timeout=3.0)
    except Exception:  # noqa: BLE001 — нет сети/сбой → пауза 5 минут
        _gt_down_until = _time.time() + 300
        return text
    _findings_ru_cache[t] = ru
    return ru


def _report_ru(rep_dict: dict) -> dict:
    """Перевести человекочитаемые поля замечаний (problem/judge_reason) на русский
    для UI. quote/locator/block/responsible_node не трогаем (нужны как есть)."""
    for f in rep_dict.get("findings") or []:
        if f.get("problem"):
            f["problem"] = _to_ru_display(f["problem"])
        if f.get("judge_reason"):
            f["judge_reason"] = _to_ru_display(f["judge_reason"])
    return rep_dict


def _serialize(values: dict) -> dict:
    """State → JSON-safe dict. findings получают актуальный статус/коммент из decisions."""
    decisions = values.get("finding_decisions") or {}
    out: dict[str, Any] = {}
    for key, val in values.items():
        if key == "chapters":
            out[key] = [c.model_dump() for c in val]
        elif key == "editor_reports":
            out[key] = [_report_ru(apply_decisions(r, decisions).model_dump())
                        for r in val]
        else:
            out[key] = val
    return out


def _run_worker(run: Run, _graph_unused, stream_input):
    """Гоняем step-граф (interrupt после каждой стадии) в ЦИКЛЕ. На каждой
    паузе-точке смотрим run.step_mode: ON → стоп и ждём resume; OFF → сами
    продолжаем. Так тумблер пошагового режима работает НА ЛЕТУ."""
    # live-стрим: узлы графа шлют нарастающий текст сюда → во фронт (poll/SSE)
    def _on_delta(stage, idx, text):
        if run.cancel:  # «Стоп» во время стрима → прерываем LLM немедленно
            raise _Cancelled()
        run.gen = {"stage": stage, "idx": idx, "text": text}
        run.events.put({"type": "gen", "stage": stage, "idx": idx, "text": text})
    cfg = {"configurable": {"thread_id": run.thread_id, "on_delta": _on_delta}}
    graph = _GRAPH_STEP
    run.cancel = False
    run.status = "running"
    run.error = ""  # новый прогон → чистим прошлую ошибку
    run.events.put({"type": "status", "status": "running"})
    inp = stream_input
    try:
        while True:
            for event in graph.stream(inp, cfg, stream_mode="values"):
                run.gen = None  # узел завершился → артефакт в state, гасим партиал
                log = event.get("log", [])
                for line in log[run.printed:]:
                    run.events.put({"type": "log", "line": line})
                run.printed = len(log)
                run.events.put({"type": "state", "state": _serialize(event)})

            run.gen = None
            snap = graph.get_state(cfg)
            if not snap.next:  # END
                run.status = "done"
                run.events.put({"type": "status", "status": "done"})
                return
            if run.cancel or run.step_mode:  # «Стоп» или пошаговый → пауза, ждём юзера
                run.cancel = False
                run.status = "paused"
                run.events.put({"type": "status", "status": "paused",
                                "next": list(snap.next)})
                return
            inp = None  # авто-режим → продолжаем со следующей стадии
    except _Cancelled:  # «Стоп» посреди этапа → встаём на паузу, прогресс цел
        run.gen = None
        run.cancel = False
        run.status = "paused"
        snap = graph.get_state(cfg)
        run.events.put({"type": "status", "status": "paused",
                        "next": list(snap.next)})
        return
    except LLMLimitError as exc:  # #5: лимит провайдера исчерпан → спросить юзера
        run.gen = None
        limit = {"provider": exc.provider, "reset_at": exc.reset_at,
                 "kind": exc.kind, "message": str(exc)[:200]}
        run.limit = limit
        run.status = "limit"
        try:
            _patch(run.thread_id, run, {"limit_info": limit})
        except Exception:  # noqa: BLE001
            pass
        run.events.put({"type": "status", "status": "limit", "limit": limit})
    except Exception as exc:  # noqa: BLE001
        run.gen = None
        run.status = "error"
        run.error = str(exc)
        run.events.put({"type": "status", "status": "error", "error": str(exc)})


def _busy(run: Run):
    if run.worker and run.worker.is_alive():
        raise HTTPException(409, "run is busy")


def _patch(thread_id: str, run: Run, patch: dict, as_node: str | None = None):
    with _LOCK:
        if as_node:
            _graph(run).update_state(_config(thread_id), patch, as_node=as_node)
        else:
            _graph(run).update_state(_config(thread_id), patch)


# ---------- модели запросов ----------

class StartReq(BaseModel):
    theme: str
    genre: Optional[str] = None
    target_language: str = "English"
    chapters_per_batch: int = 3
    translation_enabled: bool = True  # Бот 8: перевод на все языки (Google)
    step_mode: bool = True
    chapter_model: Optional[str] = None  # модель для глав (Бот 5, не-адалт)
    default_words: int = 0               # цель слов/главу (0 → дефолт конфига)
    prompt_overrides: Optional[dict] = None  # dev-оверрайды промптов до старта


class PatchReq(BaseModel):
    patch: dict[str, Any]


class ChatReq(BaseModel):
    """Чат с ИИ-редактором (#1): по главе или по конкретному замечанию."""
    messages: list[dict]                 # [{role: user|assistant, content}]
    chapter_idx: Optional[int] = None    # контекст: глава
    finding_id: Optional[str] = None     # контекст: замечание


class ReviseReq(BaseModel):
    feedback: str
    chapter_idx: Optional[int] = None


class SelectLoglineReq(BaseModel):
    logline: str


class FindingReq(BaseModel):
    status: Optional[str] = None   # "accepted" | "rejected" | "open"
    comment: Optional[str] = None
    judge: bool = False            # True → ИИ сам решает принять/отклонить


class RollbackReq(BaseModel):
    stage: str


# ---------- основные эндпоинты ----------

@app.post("/api/runs")
def start_run(req: StartReq):
    thread_id = uuid.uuid4().hex[:12]
    run = Run(thread_id=thread_id, step_mode=req.step_mode)
    RUNS[thread_id] = run
    init = {
        "theme": req.theme,
        "genre": req.genre,
        "target_language": req.target_language,
        "chapters_per_batch": max(1, req.chapters_per_batch),
        "translation_enabled": req.translation_enabled,
        "chapter_model": req.chapter_model,
        "default_words": req.default_words or CHAPTER_WORDS_DEFAULT,
    }
    # dev-оверрайды промптов, заданные ДО старта (применяются с первого бота)
    if req.prompt_overrides:
        init["prompt_overrides"] = {
            k: v for k, v in req.prompt_overrides.items() if (v or "").strip()
        }
    run.worker = threading.Thread(
        target=_run_worker, args=(run, _graph(run), init), daemon=True)
    run.worker.start()
    return {"thread_id": thread_id, "step_mode": req.step_mode}


@app.get("/api/runs/{thread_id}/state")
def get_state(thread_id: str):
    run = _get_run(thread_id)
    snap = _graph(run).get_state(_config(thread_id))
    return {"status": run.status, "next": list(snap.next),
            "error": run.error, "limit": run.limit, "gen": run.gen,
            "structure_dirty": run.structure_dirty,
            "state": _serialize(snap.values or {})}


@app.get("/api/runs")
def list_runs():
    """Список всех прогонов (новелл) из чекпоинт-БД — чтобы продолжить
    оборванную работу после рестарта/краша или скачать готовую."""
    import sqlite3
    db = os.environ.get("NARRATIVE_DB", "narrative_state.db")
    try:
        conn = sqlite3.connect(f"file:{db}?mode=ro", uri=True,
                               check_same_thread=False)
        rows = [r[0] for r in conn.execute(
            "SELECT thread_id FROM checkpoints GROUP BY thread_id "
            "ORDER BY MAX(rowid) DESC LIMIT 100")]
        conn.close()
    except Exception:
        rows = list(RUNS.keys())
    out = []
    for tid in rows:
        try:
            snap = _GRAPH_STEP.get_state(_config(tid))
        except Exception:
            continue
        vals = snap.values or {}
        if not vals.get("theme"):
            continue  # пустой/мусорный чекпоинт — пропускаем
        chs = vals.get("chapters") or []
        run = RUNS.get(tid)
        status = run.status if run else ("done" if not snap.next else "paused")
        out.append({
            "thread_id": tid,
            "theme": str(vals.get("theme", ""))[:100],
            "chapters": len(chs),
            "written": sum(1 for c in chs if getattr(c, "dialogue", None)),
            "status": status,
        })
    return {"runs": out}


@app.patch("/api/runs/{thread_id}/state")
def patch_state(thread_id: str, req: PatchReq):
    run = _get_run(thread_id)
    _busy(run)
    patch = dict(req.patch)
    if "chapters" in patch:
        patch["chapters"] = [Chapter(**c) for c in patch["chapters"]]
    _patch(thread_id, run, patch)
    run.events.put({"type": "edited", "keys": list(patch.keys())})
    return {"ok": True}


@app.post("/api/runs/{thread_id}/resume")
def resume_run(thread_id: str):
    run = _get_run(thread_id)
    _busy(run)
    run.limit = None  # снимаем баннер лимита при ресюме
    with _LOCK:  # TOCTOU: двойной resume → два воркера на один граф
        if run.worker and run.worker.is_alive():
            raise HTTPException(409, "run is busy")
        run.worker = threading.Thread(
            target=_run_worker, args=(run, None, None), daemon=True)
        run.worker.start()
    return {"ok": True}


@app.post("/api/runs/{thread_id}/stop")
def stop_run(thread_id: str):
    """Остановить генерацию на ЛЮБОМ этапе: ставим флаг — воркер мягко встаёт на
    паузу (на стрим-этапах сразу, на остальных — по завершении текущего бота).
    Прогресс сохранён, можно продолжить кнопкой «Продолжить»."""
    run = _get_run(thread_id)
    if run.worker and run.worker.is_alive():
        run.cancel = True
    elif run.status == "running":  # завис без воркера — снимаем; done/error не трогаем
        run.status = "paused"
        run.events.put({"type": "status", "status": "paused"})
    return {"ok": True, "stopping": True}


# ---------- #5: провайдер по этапам + разрешение лимита ----------

_VALID_STAGES = set(STAGE_NODES) | {"locations", "chat"}


class StageProviderReq(BaseModel):
    stage: str                  # этап из STAGE_NODES (или "all")
    provider: str               # "subscription" | "openrouter" | "default"


@app.post("/api/runs/{thread_id}/stage_provider")
def set_stage_provider(thread_id: str, req: StageProviderReq):
    """Выбор провайдера (подписка/OpenRouter) для конкретного этапа (#5)."""
    run = _get_run(thread_id)
    if req.provider not in ("subscription", "openrouter", "default"):
        raise HTTPException(400, "bad provider")
    st = _state(thread_id)
    sp = dict(st.get("stage_providers") or {})
    stages = list(_VALID_STAGES) if req.stage == "all" else [req.stage]
    for s in stages:
        if req.provider == "default":
            sp.pop(s, None)
        else:
            sp[s] = req.provider
    _patch(thread_id, run, {"stage_providers": sp})
    return {"ok": True, "stage_providers": sp}


class LimitResolveReq(BaseModel):
    action: str                 # "switch" (→OpenRouter) | "wait" | "subscription"


@app.post("/api/runs/{thread_id}/limit/resolve")
def resolve_limit(thread_id: str, req: LimitResolveReq):
    """Решение нарративщика по исчерпанному лимиту (#5).

    switch — весь ран на OpenRouter и продолжить; subscription — снова на
    подписку (после сброса) и продолжить; wait — просто снять баннер.
    """
    run = _get_run(thread_id)
    _busy(run)
    if req.action == "switch":
        _patch(thread_id, run, {"force_openrouter": True, "limit_info": None})
    elif req.action == "subscription":
        _patch(thread_id, run, {"force_openrouter": False, "limit_info": None})
    elif req.action != "wait":
        raise HTTPException(400, "bad action")
    run.limit = None
    if req.action in ("switch", "subscription"):
        run.worker = threading.Thread(
            target=_run_worker, args=(run, None, None), daemon=True)
        run.worker.start()
        return {"ok": True, "resumed": True}
    return {"ok": True, "resumed": False}


class StepReq(BaseModel):
    enabled: bool


@app.post("/api/runs/{thread_id}/step")
def set_step(thread_id: str, req: StepReq):
    """Тумблер пошагового режима НА ЛЕТУ. Если выключили во время паузы —
    автоматически продолжаем (worker добежит до конца / след. паузы)."""
    run = _get_run(thread_id)
    auto_continued = False
    # под _LOCK: иначе TOCTOU между проверкой run.status/worker и стартом потока
    # (воркер параллельно читает step_mode / пишет status) → 2 воркера разом.
    with _LOCK:
        run.step_mode = req.enabled
        if not req.enabled and run.status == "paused" and not (
                run.worker and run.worker.is_alive()):
            run.worker = threading.Thread(
                target=_run_worker, args=(run, None, None), daemon=True)
            run.worker.start()
            auto_continued = True
    return {"ok": True, "step_mode": run.step_mode, "resumed": auto_continued}


# ---------- Бот 1: выбор логлайна ----------

@app.post("/api/runs/{thread_id}/select_logline")
def select_logline(thread_id: str, req: SelectLoglineReq):
    run = _get_run(thread_id)
    _busy(run)
    _patch(thread_id, run, {"selected_logline": req.logline})
    run.events.put({"type": "edited", "keys": ["selected_logline"]})
    return {"ok": True}


# ---------- per-stage revise: попросить ИИ переделать этап по фидбеку ----------

@app.post("/api/runs/{thread_id}/stage/{stage}/revise")
def revise_stage(thread_id: str, stage: str, req: ReviseReq):
    """Попросить ИИ переделать этап (ФОНОВО — LLM долгий)."""
    run = _get_run(thread_id)
    _busy(run)
    if stage not in NODE_FUNCS:
        raise HTTPException(400, f"unknown stage {stage}")
    # поглавные узлы требуют chapter_idx (из запроса или уже в state) — иначе
    # node упадёт на state["chapter_idx"] с KeyError → 500 с трейсом.
    if stage in ("dialogue", "adult", "editor"):
        st0 = _state(thread_id)
        if req.chapter_idx is None and st0.get("chapter_idx") is None:
            raise HTTPException(400, f"stage {stage} требует chapter_idx")

    def _op():
        st = _state(thread_id)
        st["revision_feedback"] = nodes.to_english(req.feedback)  # RU→EN для агента
        if stage in ("dialogue", "adult"):
            st["revision_target"] = stage
        if req.chapter_idx is not None:
            st["chapter_idx"] = req.chapter_idx
        result = NODE_FUNCS[stage](st)
        _patch(thread_id, run, result)
        run.events.put({"type": "edited", "keys": list(result.keys())})

    return _bg_op(run, _op, stage=stage, idx=req.chapter_idx)


# ---------- правка плана конкретной главы (структура) ----------

@app.post("/api/runs/{thread_id}/chapter/{idx}/revise")
def revise_chapter(thread_id: str, idx: int, req: ReviseReq):
    """Переписать план ОДНОЙ главы по фидбеку (ФОНОВО)."""
    run = _get_run(thread_id)
    _busy(run)
    st = _state(thread_id)
    if idx < 0 or idx >= len(st.get("chapters") or []):
        raise HTTPException(400, "bad chapter index")

    def _op():
        from narrative import prompts  # noqa: PLC0415
        st2 = _state(thread_id)
        chapters = list(st2.get("chapters") or [])
        ch = chapters[idx]
        loc = (st2.get("locations") or "").strip()
        user = (
            f"Карточки:\n{st2.get('characters','')}\n\n"
            + (f"Локации:\n{loc}\n\n" if loc else "")
            + nodes._story_context(st2, idx).strip()
            + f"\n\nГлава {idx + 1}: {ch.title}\nТекущий план:\n{ch.plan}\n\n"
            f"Правка нарративщика: {nodes.to_english(req.feedback)}\n\n"
            "Перепиши план ТОЛЬКО этой главы с учётом правки, сохранив формат "
            "разделов (LOCATIONS/CHARACTERS/STORY/CHOICE/ADULT SCENES) и "
            "непрерывность с остальными. Верни только новый текст плана, "
            "без преамбул."
        )
        ch.plan = nodes._structural(st2, "structure").complete(
            prompts.with_genre(prompts.STRUCTURE, st2.get("genre"))
            + prompts.NAMES_RULE, user)
        chapters[idx] = ch
        _patch(thread_id, run, {"chapters": chapters})
        run.events.put({"type": "edited", "keys": ["chapters"]})

    return _bg_op(run, _op)


@app.post("/api/runs/{thread_id}/chapter/{idx}/rewrite_dialogue")
def rewrite_dialogue(thread_id: str, idx: int):
    """Правка плана → переписать диалоги (и адалт) главы (ФОНОВО)."""
    run = _get_run(thread_id)
    _busy(run)
    st = _state(thread_id)
    if idx < 0 or idx >= len(st.get("chapters") or []):
        raise HTTPException(400, "bad chapter index")

    def _op():
        st2 = _state(thread_id)
        chapters = list(st2.get("chapters") or [])
        chapters[idx] = nodes.write_chapter(st2, chapters[idx], idx)
        _patch(thread_id, run, {"chapters": chapters})
        run.events.put({"type": "edited", "keys": ["chapters"]})

    return _bg_op(run, _op)


@app.post("/api/runs/{thread_id}/chapter/{idx}/sync_plan")
def sync_plan(thread_id: str, idx: int):
    """Правка диалогов → подогнать план главы под текст (ФОНОВО)."""
    run = _get_run(thread_id)
    _busy(run)
    st = _state(thread_id)
    if idx < 0 or idx >= len(st.get("chapters") or []):
        raise HTTPException(400, "bad chapter index")

    def _op():
        st2 = _state(thread_id)
        chapters = list(st2.get("chapters") or [])
        ch = chapters[idx]
        ch.plan = nodes.sync_plan_from_dialogue(st2, ch, idx)
        chapters[idx] = ch
        _patch(thread_id, run, {"chapters": chapters})
        run.events.put({"type": "edited", "keys": ["chapters"]})

    return _bg_op(run, _op)


# ---------- адалт: адаптировать главу / отключить ----------

@app.post("/api/runs/{thread_id}/chapter/{idx}/adapt_adult")
def adapt_adult(thread_id: str, idx: int):
    """Глава без почвы для адалта → Бот 5 вплетает подводку, Бот 6 генерит
    сцену (ФОНОВО — два долгих LLM-вызова, grok)."""
    run = _get_run(thread_id)
    _busy(run)
    st = _state(thread_id)
    if idx < 0 or idx >= len(st.get("chapters") or []):
        raise HTTPException(400, "bad chapter index")

    def _op():
        st2 = _state(thread_id)
        chapters = list(st2.get("chapters") or [])
        ch = chapters[idx]
        bridge = ch.adult_bridge_hint or (
            "добавь взаимное влечение и эротическое напряжение между подходящей "
            "парой персонажей и повод остаться наедине к финалу главы"
        )
        st2["chapter_idx"] = idx
        st2["revision_target"] = "dialogue"
        st2["revision_feedback"] = (
            f"Адаптируй главу под откровенную сцену: {bridge}. "
            "Сохрани сюжет и события главы, не переписывай всё — органично "
            "вплети подводку так, чтобы финал естественно переходил в адалт."
        )
        res_dlg = nodes.dialogue_node(st2)
        st2.update(res_dlg)
        st2["chapter_idx"] = idx
        res_adult = nodes.adult_node(st2)
        patch = {**res_dlg, **res_adult}
        patch["log"] = (res_dlg.get("log") or []) + (res_adult.get("log") or [])
        _patch(thread_id, run, patch)
        run.events.put({"type": "edited", "keys": ["chapters"]})

    return _bg_op(run, _op)


@app.post("/api/runs/{thread_id}/chapter/{idx}/skip_adult")
def skip_adult(thread_id: str, idx: int):
    """Оставить главу без адалта: снять адалт-точку и блокировку."""
    run = _get_run(thread_id)
    _busy(run)
    st = _state(thread_id)
    chapters = list(st.get("chapters") or [])
    if idx < 0 or idx >= len(chapters):
        raise HTTPException(400, "bad chapter index")
    chapters[idx].is_adult_point = False
    chapters[idx].adult_block_reason = None
    chapters[idx].adult_bridge_hint = None
    _patch(thread_id, run, {"chapters": chapters})
    run.events.put({"type": "edited", "keys": ["chapters"]})
    return {"ok": True}


# ---------- структура порциями ----------

class CountReq(BaseModel):
    count: int


@app.post("/api/runs/{thread_id}/chapter_count")
def set_chapter_count(thread_id: str, req: CountReq):
    """Утвердить число глав (апрув предложения ИИ или своё) → продолжить к структуре."""
    run = _get_run(thread_id)
    _busy(run)
    n = max(2, min(int(req.count), 30))
    _patch(thread_id, run, {"target_chapters": n})
    return resume_run(thread_id)


@app.post("/api/runs/{thread_id}/restructure")
def restructure(thread_id: str, req: CountReq):
    """#7: пересобрать структуру под НОВОЕ число глав (растянуть/сжать сюжет,
    не обрезать). Регенерит поглавный план — написанные тексты сбрасываются."""
    run = _get_run(thread_id)
    _busy(run)
    n = max(2, min(int(req.count), 30))

    def _op():
        st = _state(thread_id)
        st["target_chapters"] = n
        st["manual_chapters"] = False  # явная пересборка → снимаем ручной режим
        result = nodes.structure_node(st)  # учитывает текущий план как опору
        result["manual_chapters"] = False
        _patch(thread_id, run, result)
        run.events.put({"type": "edited", "keys": list(result.keys())})

    return _bg_op(run, _op)


class AddChapterReq(BaseModel):
    after_idx: int = -1     # вставить ПОСЛЕ этого индекса; -1 → в начало
    is_adult: bool = True   # адалт-глава (по умолчанию да) или обычная
    generate: bool = False  # True → ИИ сразу пишет план; False → пустой шаблон


@app.post("/api/runs/{thread_id}/chapter/add")
def add_chapter(thread_id: str, req: AddChapterReq):
    """Добавить ОДНУ главу. По умолчанию — ПУСТОЙ черновик (без ИИ): план/текст
    пишутся вручную или генерятся отдельной кнопкой. generate=True → ИИ сразу
    напишет план (связав соседей)."""
    run = _get_run(thread_id)
    _busy(run)

    def _insert_patch(st: dict, chapters: list, pos: int) -> dict:
        """Патч вставки: главы + сдвиг курсора, если вставили ДО него."""
        patch: dict[str, Any] = {"chapters": chapters,
                                 "revision_target": None,
                                 "revision_feedback": None}
        rc = _remap_retry(st, lambda n: n + 1 if n >= pos else n)
        if rc is not None:
            patch["retry_count"] = rc
        cur = st.get("chapter_idx", 0)
        if pos <= cur:
            patch["chapter_idx"] = cur + 1
        return patch

    if not req.generate:
        # пустой шаблон — синхронно, без LLM
        st = _state(thread_id)
        chapters = list(st.get("chapters") or [])
        after = max(-1, min(req.after_idx, len(chapters) - 1))
        new_ch = Chapter(index=after + 1, title="Новая глава", plan="",
                         is_adult_point=req.is_adult, adult_note="")
        chapters.insert(after + 1, new_ch)
        for i, ch in enumerate(chapters):
            ch.index = i
        _patch(thread_id, run, _insert_patch(st, chapters, after + 1))
        run.structure_dirty = True
        run.events.put({"type": "edited", "keys": ["chapters"]})
        return {"ok": True, "chapters": len(chapters), "new_index": after + 1}

    def _op():
        st = _state(thread_id)
        chapters = list(st.get("chapters") or [])
        after = max(-1, min(req.after_idx, len(chapters) - 1))
        new_ch = nodes.gen_inserted_chapter(st, after)
        new_ch.is_adult_point = req.is_adult  # адалт/неадалт по выбору юзера
        if not req.is_adult:
            new_ch.adult_note = ""
        chapters.insert(after + 1, new_ch)
        for i, ch in enumerate(chapters):
            ch.index = i
        _patch(thread_id, run, _insert_patch(st, chapters, after + 1))
        run.structure_dirty = True
        run.events.put({"type": "edited", "keys": ["chapters"]})

    return _bg_op(run, _op)


@app.post("/api/runs/{thread_id}/chapter/{idx}/gen_plan")
def gen_chapter_plan(thread_id: str, idx: int):
    """Сгенерировать ИИ план для (обычно пустой) главы idx — видит все главы."""
    run = _get_run(thread_id)
    _busy(run)
    # валидация ДО старта потока — иначе клиент получал «ok» и лишь потом ошибку
    if idx < 0 or idx >= len(_state(thread_id).get("chapters") or []):
        raise HTTPException(400, "bad chapter index")

    def _op():
        st = _state(thread_id)
        chapters = list(st.get("chapters") or [])
        if idx >= len(chapters):  # главу успели удалить — тихий no-op
            return
        chapters[idx].plan = nodes.gen_chapter_plan(st, idx)
        _patch(thread_id, run, {"chapters": chapters})
        run.events.put({"type": "edited", "keys": ["chapters"]})

    return _bg_op(run, _op)


@app.post("/api/runs/{thread_id}/structure/manual")
def manual_structure(thread_id: str):
    """Пропустить выбор объёма и ИИ-структуру: создать ОДИН пустой черновик главы.
    Дальше нарративщик добавляет главы и пишет/генерит планы вручную."""
    run = _get_run(thread_id)
    _busy(run)
    ch = Chapter(index=0, title="Глава 1", plan="", is_adult_point=True,
                 adult_note="")
    patch = {"chapters": [ch], "target_chapters": 1, "structure_done": True,
             "structure_fixes": [], "phase": "content", "chapter_idx": 0,
             "manual_chapters": True}
    # manual_chapters=True → даже если граф вернётся на structure/structure_editor
    # (напр. после ручной правки глав), эти узлы отработают вхолостую и НЕ затрут
    # ручные главы. as_node="structure_editor" сразу ведёт к написанию.
    try:
        _patch(thread_id, run, patch, as_node="structure_editor")
    except Exception:  # noqa: BLE001 — фолбэк, если граф не там
        _patch(thread_id, run, patch)
    run.structure_dirty = False
    run.status = "paused"
    run.error = ""
    run.events.put({"type": "edited", "keys": ["chapters"]})
    run.events.put({"type": "status", "status": "paused"})
    return {"ok": True}


def _remap_retry(st: dict, remap) -> dict | None:
    """Пересчитать ключи retry_count ("node:chapterIdx") при переиндексации глав.
    remap(old_idx) -> new_idx | None (None = глава удалена, счётчик отбрасываем)."""
    rc = st.get("retry_count") or {}
    if not rc:
        return None
    out: dict[str, int] = {}
    for key, cnt in rc.items():
        node, _, sidx = key.rpartition(":")
        try:
            old = int(sidx)
        except ValueError:
            out[key] = cnt
            continue
        new = remap(old)
        if new is not None:
            out[f"{node}:{new}"] = cnt
    return out


@app.delete("/api/runs/{thread_id}/chapter/{idx}")
def delete_chapter(thread_id: str, idx: int):
    """#7: удалить главу и переиндексировать остальные."""
    run = _get_run(thread_id)
    _busy(run)
    st = _state(thread_id)
    chapters = list(st.get("chapters") or [])
    if idx < 0 or idx >= len(chapters):
        raise HTTPException(400, "bad chapter index")
    if len(chapters) <= 1:
        raise HTTPException(400, "нельзя удалить последнюю главу")
    chapters.pop(idx)
    for i, ch in enumerate(chapters):
        ch.index = i
    # курсор «какую главу пишем следующей» должен указывать на ТУ ЖЕ главу:
    # удалили главу ДО курсора → курсор сдвигается на -1 (иначе пропуск главы)
    cur = st.get("chapter_idx", 0)
    new_cur = cur - 1 if idx < cur else cur
    new_cur = max(0, min(new_cur, len(chapters) - 1))
    # незавершённая ревизия привязана к индексу — после сдвига попала бы в чужую
    # главу, поэтому сбрасываем
    patch = {"chapters": chapters,
             "revision_target": None, "revision_feedback": None}
    rc = _remap_retry(st, lambda n: None if n == idx else (n - 1 if n > idx else n))
    if rc is not None:
        patch["retry_count"] = rc
    if new_cur != cur:
        patch["chapter_idx"] = new_cur
    _patch(thread_id, run, patch)
    run.structure_dirty = True
    run.events.put({"type": "edited", "keys": ["chapters"]})
    return {"ok": True, "chapters": len(chapters)}


class ReorderReq(BaseModel):
    order: list[int]          # новая последовательность СТАРЫХ индексов глав


@app.post("/api/runs/{thread_id}/chapter/reorder")
def reorder_chapters(thread_id: str, req: ReorderReq):
    """Drag-and-drop: задать новый порядок глав целиком. order — перестановка
    текущих индексов (порядок в массиве = порядок повествования = порядок в коде)."""
    run = _get_run(thread_id)
    _busy(run)
    st = _state(thread_id)
    chapters = list(st.get("chapters") or [])
    if sorted(req.order) != list(range(len(chapters))):
        raise HTTPException(400, "bad order permutation")
    new = [chapters[i] for i in req.order]
    for i, ch in enumerate(new):
        ch.index = i
    patch: dict[str, Any] = {"chapters": new,
                             "revision_target": None,
                             "revision_feedback": None}
    rc = _remap_retry(st, lambda n: req.order.index(n) if n in req.order else None)
    if rc is not None:
        patch["retry_count"] = rc
    # курсор следует за СВОЕЙ главой на её новое место
    cur = st.get("chapter_idx", 0)
    if 0 <= cur < len(req.order):
        new_cur = req.order.index(cur)
        if new_cur != cur:
            patch["chapter_idx"] = new_cur
    _patch(thread_id, run, patch)
    run.events.put({"type": "edited", "keys": ["chapters"]})
    return {"ok": True}


class MoveReq(BaseModel):
    dir: str                  # "up" | "down"


@app.post("/api/runs/{thread_id}/chapter/{idx}/move")
def move_chapter(thread_id: str, idx: int, req: MoveReq):
    """#6: двигать главу вверх/вниз — меняет порядок и переиндексирует (порядок
    в коде = порядок в массиве)."""
    run = _get_run(thread_id)
    _busy(run)
    st = _state(thread_id)
    chapters = list(st.get("chapters") or [])
    if idx < 0 or idx >= len(chapters):
        raise HTTPException(400, "bad chapter index")
    j = idx - 1 if req.dir == "up" else idx + 1
    if j < 0 or j >= len(chapters):
        return {"ok": True, "chapters": len(chapters)}  # край — ничего
    chapters[idx], chapters[j] = chapters[j], chapters[idx]
    for i, ch in enumerate(chapters):
        ch.index = i
    patch: dict[str, Any] = {"chapters": chapters,
                             "revision_target": None,
                             "revision_feedback": None}
    rc = _remap_retry(st, lambda n: j if n == idx else (idx if n == j else n))
    if rc is not None:
        patch["retry_count"] = rc
    # курсор следует за своей главой при свопе
    cur = st.get("chapter_idx", 0)
    if cur == idx:
        patch["chapter_idx"] = j
    elif cur == j:
        patch["chapter_idx"] = idx
    _patch(thread_id, run, patch)
    run.events.put({"type": "edited", "keys": ["chapters"]})
    return {"ok": True, "moved_to": j}


@app.post("/api/runs/{thread_id}/structure/check")
def check_structure(thread_id: str):
    """Проверка структуры (бот-редактор) после ручных изменений глав. Чинит
    нестыковки плана. Снимает флаг structure_dirty."""
    run = _get_run(thread_id)
    _busy(run)
    run.structure_dirty = False

    def _op():
        st = _state(thread_id)
        result = nodes.structure_editor_node(st)
        _patch(thread_id, run, result)
        run.events.put({"type": "edited", "keys": list(result.keys())})

    return _bg_op(run, _op, stage="structure_editor")


@app.post("/api/runs/{thread_id}/structure/skip_check")
def skip_structure_check(thread_id: str):
    """Пропустить проверку структуры — просто снять флаг."""
    run = _get_run(thread_id)
    run.structure_dirty = False
    run.events.put({"type": "status", "status": run.status})
    return {"ok": True}


@app.post("/api/runs/{thread_id}/structure/skip_stage")
def skip_structure_stage(thread_id: str):
    """Пропустить ЭТАП проверки структуры в пайплайне (узел structure_editor) —
    не запускать редактора, сразу перейти к написанию глав."""
    run = _get_run(thread_id)
    _busy(run)
    run.structure_dirty = False
    # помечаем узел выполненным без правок → граф идёт к следующему шагу (главы).
    # Фолбэк: если граф не на узле проверки — ставим одноразовый флаг пропуска,
    # тогда узел отработает вхолостую (без 500 и без перезаписи плана).
    try:
        _patch(thread_id, run, {"structure_fixes": [],
                                "skip_structure_check": True},
               as_node="structure_editor")
    except Exception:  # noqa: BLE001
        _patch(thread_id, run, {"skip_structure_check": True})
    return resume_run(thread_id)


# ---------- объём главы (слова) + «растянуть» ----------

class WordsReq(BaseModel):
    words: int                # цель слов на главу (0 → дефолт прогона)


@app.post("/api/runs/{thread_id}/chapter/{idx}/words")
def set_chapter_words(thread_id: str, idx: int, req: WordsReq):
    """Задать целевой объём (слова) для конкретной главы (оверрайд дефолта)."""
    run = _get_run(thread_id)
    _busy(run)
    st = _state(thread_id)
    chapters = list(st.get("chapters") or [])
    if idx < 0 or idx >= len(chapters):
        raise HTTPException(400, "bad chapter index")
    w = int(req.words)
    chapters[idx].target_words = w if w > 0 else None
    _patch(thread_id, run, {"chapters": chapters})
    run.events.put({"type": "edited", "keys": ["chapters"]})
    return {"ok": True, "target_words": chapters[idx].target_words}


class ExpandReq(BaseModel):
    add_words: int = 800      # на сколько примерно растянуть


@app.post("/api/runs/{thread_id}/chapter/{idx}/expand")
def expand_chapter_ep(thread_id: str, idx: int, req: ExpandReq):
    """«Растянуть» текст главы на ~add_words (ФОНОВО — LLM долгий)."""
    run = _get_run(thread_id)
    _busy(run)
    st = _state(thread_id)
    chapters = list(st.get("chapters") or [])
    if idx < 0 or idx >= len(chapters):
        raise HTTPException(400, "bad chapter index")

    def _op():
        st2 = _state(thread_id)
        chs = list(st2.get("chapters") or [])
        ch = nodes.expand_chapter(st2, chs[idx], idx, max(100, req.add_words))
        chs[idx] = ch
        _patch(thread_id, run, {"chapters": chs})
        run.events.put({"type": "edited", "keys": ["chapters"]})

    return _bg_op(run, _op, stage="dialogue", idx=idx)


class ExpandStageReq(BaseModel):
    add_words: int = 400


@app.post("/api/runs/{thread_id}/stage/{stage}/expand")
def expand_stage(thread_id: str, stage: str, req: ExpandStageReq):
    """«Растянуть» текстовый блок-этап (synopsis/characters/locations)."""
    run = _get_run(thread_id)
    _busy(run)
    if stage not in ("synopsis", "characters", "locations"):
        raise HTTPException(400, "expand доступен для synopsis/characters/locations")

    def _op():
        from narrative import prompts  # noqa: PLC0415
        st2 = _state(thread_id)
        cur = st2.get(stage) or ""
        if not cur.strip():
            return
        system = nodes._sys(st2, prompts.STRUCTURE, names=True)
        user = (f"ТЕКУЩИЙ ТЕКСТ:\n{cur}"
                + prompts.EXPAND_GUIDE.format(add=max(100, req.add_words)))
        out = nodes._structural(st2, stage).complete(system, user)
        if out.strip() and len(out) > len(cur):
            _patch(thread_id, run, {stage: out})
            run.events.put({"type": "edited", "keys": [stage]})

    return _bg_op(run, _op, stage=stage)


# ---------- решения по findings редактора ----------

def _find_finding(st: dict, fid: str):
    for r in st.get("editor_reports") or []:
        for f in r.findings:
            if f.id == fid:
                return f
    return None


# Инструкция «без markdown»: чат-ответы ИИ должны быть простым текстом.
NO_MARKDOWN = (
    " ВАЖНО: отвечай ПРОСТЫМ текстом без markdown-разметки — не используй "
    "звёздочки (* **), решётки (#), обратные кавычки/блоки кода (```), "
    "маркированные/нумерованные списки и таблицы. Только обычные абзацы."
)


def _last_report(st: dict, idx: int):
    """Последний отчёт редактора по главе idx (с наложенными решениями) или None."""
    decisions = st.get("finding_decisions") or {}
    reports = [r for r in (st.get("editor_reports") or [])
               if r.chapter_index == idx]
    if not reports:
        return None
    return apply_decisions(reports[-1], decisions)


@app.post("/api/runs/{thread_id}/findings/{fid}")
def decide_finding(thread_id: str, fid: str, req: FindingReq):
    run = _get_run(thread_id)
    st = _state(thread_id)
    decisions = dict(st.get("finding_decisions") or {})
    d = dict(decisions.get(fid, {}))
    if req.comment is not None:
        d["comment"] = req.comment
    if req.status:
        d["status"] = req.status
    if req.judge:  # ИИ сам решает — strict JSON, вердикт с обоснованием
        f = _find_finding(st, fid)
        if f is not None:
            verdict = nodes._editor().structured(
                "Ты — главный редактор визуальной новеллы. Реши, СПРАВЕДЛИВО ли "
                "замечание и нужно ли его исправлять. Учитывай комментарий "
                "нарративщика, если он есть.",
                f"Замечание [{f.severity}/{f.block}]: {f.problem}\n"
                f"Фрагмент: «{f.quote}»\n"
                f"Комментарий нарративщика: {d.get('comment','') or '—'}",
                JudgeOut,
            )
            d["status"] = "rejected" if verdict.decision == "reject" else "accepted"
            d["judge_reason"] = verdict.reason
            d["judged"] = True
    decisions[fid] = d
    patch: dict[str, Any] = {"finding_decisions": decisions}
    # #8: отклонённое — запоминаем текст, чтобы редактор больше не поднимал
    if d.get("status") == "rejected":
        f = _find_finding(st, fid)
        if f is not None:
            notes = list(st.get("rejected_notes") or [])
            note = f.problem.strip()
            if note and note not in notes:
                notes.append(note)
                patch["rejected_notes"] = notes
    _patch(thread_id, run, patch)
    run.events.put({"type": "finding", "id": fid, "decision": d})
    return {"ok": True, "decision": d, "state": _serialize(_state(thread_id))}


# ---------- чат с ИИ-редактором (#1) ----------

@app.post("/api/runs/{thread_id}/chat")
def chat(thread_id: str, req: ChatReq):
    """Обсудить правку с ИИ: контекст — глава целиком или конкретное замечание.

    Не меняет state — только разговор. Применение правок нарративщик делает
    обычными инструментами (правка текста / «попросить ИИ» / принять finding).
    """
    _get_run(thread_id)
    st = _state(thread_id)
    chs = list(st.get("chapters") or [])

    ctx_parts = [f"Карточки персонажей:\n{st.get('characters', '')}"]
    if st.get("locations"):
        ctx_parts.append(f"Локации:\n{st.get('locations')}")
    if st.get("synopsis"):
        ctx_parts.append(f"Синопсис:\n{st.get('synopsis')}")

    # #4: ИИ видит планы И тексты ВСЕХ глав (даже пустых), текущая помечена —
    # чтобы не путал, о какой главе речь, и держал непрерывность повествования.
    if chs:
        ctx_parts.append(nodes._all_chapters_context(st, req.chapter_idx))

    if req.chapter_idx is not None and 0 <= req.chapter_idx < len(chs):
        last = _last_report(st, req.chapter_idx)
        if last and last.findings:
            joined = "\n".join(
                f"- [{f.severity}/{f.block}] {f.locator}: {f.problem}"
                f"{' (ОТКЛОНЕНО нарративщиком)' if f.status == 'rejected' else ''}"
                for f in last.findings
            )
            ctx_parts.append(
                "Замечания редактора по этой главе (текущая ревизия):\n" + joined
            )
    if req.finding_id:
        f = _find_finding(st, req.finding_id)
        if f is not None:
            ctx_parts.append(
                f"Обсуждаемое замечание редактора [{f.severity}/{f.block}]:\n"
                f"{f.problem}\nФрагмент: «{f.quote}»"
            )

    system = (
        "Ты — опытный редактор визуальных новелл 18+, собеседник нарративщика. "
        "Обсуждаешь ТОЛЬКО текущую главу: её ПЛАН и ТЕКСТ (диалоги/адалт). "
        "Соседние главы видишь лишь как контекст для непрерывности — НЕ предлагай "
        "и НЕ берись менять другие главы, синопсис, персонажей, локации, "
        "структуру или порядок глав. Предлагай конкретные формулировки и фиксы "
        "по плану и тексту ЭТОЙ главы, объясняй коротко. Не переписывай главу "
        "целиком без просьбы. Отвечай на языке нарративщика." + NO_MARKDOWN
        + "\n\nКОНТЕКСТ:\n" + "\n\n".join(ctx_parts)
    )
    history = [
        {"role": m.get("role", "user"), "content": str(m.get("content", ""))}
        for m in req.messages[-16:]  # окно диалога
        if m.get("role") in ("user", "assistant")
    ]
    try:
        reply = nodes._editor(st, "chat").chat(system, history)
    except LLMLimitError as exc:  # #5: лимит → 429 + инфо для баннера
        raise HTTPException(429, _limit_detail(exc))
    except Exception as exc:  # noqa: BLE001 — LLM/провайдер недоступен
        raise HTTPException(502, f"чат недоступен (LLM): {str(exc)[:160]}")
    return {"reply": reply}


def _limit_detail(exc: LLMLimitError) -> dict:
    return {"error": "limit", "provider": exc.provider,
            "reset_at": exc.reset_at, "kind": exc.kind,
            "message": str(exc)[:200]}


# ---------- ассистент по всему проекту (отдельная вкладка) ----------

def _project_context(st: dict) -> str:
    """Полный срез проекта для ассистента: логлайн, синопсис, персонажи,
    локации, все главы (планы + тексты)."""
    parts: list[str] = []
    if st.get("selected_logline"):
        parts.append(f"Логлайн (выбранный):\n{st['selected_logline']}")
    elif st.get("loglines"):
        parts.append("Логлайны (варианты):\n" + "\n".join(st["loglines"]))
    if st.get("synopsis"):
        parts.append(f"Синопсис:\n{st['synopsis']}")
    if st.get("characters"):
        parts.append(f"Персонажи:\n{st['characters']}")
    if st.get("locations"):
        parts.append(f"Локации:\n{st['locations']}")
    allc = nodes._all_chapters_context(st)
    if allc:
        parts.append(allc)
    return "\n\n".join(parts) or "(проект пуст — ещё ничего не сгенерировано)"


class ProjectChatReq(BaseModel):
    messages: list[dict]                 # [{role, content}]


@app.post("/api/runs/{thread_id}/project_chat")
def project_chat(thread_id: str, req: ProjectChatReq):
    """Свободный чат по ВСЕМУ проекту (любой этап). Только обсуждение —
    применение правок отдельной кнопкой (project_apply)."""
    _get_run(thread_id)
    st = _state(thread_id)
    system = (
        "Ты — ассистент-сценарист по визуальным новеллам 18+, помогаешь "
        "нарративщику по ВСЕМУ проекту: логлайн, синопсис, персонажи, локации, "
        "планы и тексты глав. Можешь обсуждать что угодно по проекту независимо "
        "от этапа, предлагать конкретные правки содержимого. НО ты НЕ управляешь "
        "пайплайном: не запускаешь этапы, не меняешь их порядок, не двигаешь "
        "главы и не перескакиваешь между шагами — только содержимое результатов "
        "ботов. Применение правок делает нарративщик кнопкой «Применить». "
        "Отвечай по делу на языке нарративщика." + NO_MARKDOWN
        + "\n\nПРОЕКТ:\n" + _project_context(st)
    )
    history = [
        {"role": m.get("role", "user"), "content": str(m.get("content", ""))}
        for m in req.messages[-20:]
        if m.get("role") in ("user", "assistant")
    ]
    try:
        reply = nodes._editor(st, "chat").chat(system, history)
    except LLMLimitError as exc:
        raise HTTPException(429, _limit_detail(exc))
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(502, f"ассистент недоступен (LLM): {str(exc)[:160]}")
    return {"reply": reply}


# целевое поле → (метка, что патчим)
_APPLY_TARGETS = {"synopsis", "characters", "locations", "logline",
                  "chapter_plan", "chapter_dialogue"}


class ProjectApplyReq(BaseModel):
    messages: list[dict]
    target: str                          # см. _APPLY_TARGETS
    chapter_idx: Optional[int] = None    # для chapter_*


@app.post("/api/runs/{thread_id}/project_apply")
def project_apply(thread_id: str, req: ProjectApplyReq):
    """Применить обсуждённое в ассистент-чате к ОДНОМУ результату бота. Меняет
    только содержимое выбранного поля, не трогает пайплайн/порядок."""
    from narrative import prompts  # noqa: PLC0415
    run = _get_run(thread_id)
    _busy(run)
    if req.target not in _APPLY_TARGETS:
        raise HTTPException(400, f"bad target {req.target}")
    st = _state(thread_id)
    convo = nodes.to_english("\n".join(
        f"{'Нарративщик' if m.get('role') == 'user' else 'ИИ'}: "
        f"{m.get('content', '')}"
        for m in (req.messages or []) if m.get("role") in ("user", "assistant")
    ))

    labels = {"synopsis": "СИНОПСИС", "characters": "КАРТОЧКИ ПЕРСОНАЖЕЙ",
              "locations": "КАРТОЧКИ ЛОКАЦИЙ", "logline": "ЛОГЛАЙН",
              "chapter_plan": "ПЛАН ГЛАВЫ", "chapter_dialogue": "ТЕКСТ ГЛАВЫ"}
    chapters = list(st.get("chapters") or [])
    idx = req.chapter_idx
    if req.target in ("chapter_plan", "chapter_dialogue"):
        if idx is None or not (0 <= idx < len(chapters)):
            raise HTTPException(400, "нужен корректный chapter_idx")
        cur = (chapters[idx].plan if req.target == "chapter_plan"
               else (chapters[idx].dialogue or ""))
    elif req.target == "logline":
        cur = st.get("selected_logline") or (st.get("loglines") or [""])[0]
    else:
        cur = st.get(req.target) or ""

    system = (
        f"Ты редактор-сценарист 18+. Перепиши {labels[req.target]} с учётом "
        "обсуждения ниже. Сохрани прежний формат и стиль, поменяй только то, что "
        "обсудили. Верни ТОЛЬКО новый текст, без преамбул и пояснений."
        + prompts.NAMES_RULE + prompts.CONTENT_POLICY
        + "\n\nКОНТЕКСТ ПРОЕКТА:\n" + _project_context(st)
    )
    user = (f"ТЕКУЩЕЕ содержимое ({labels[req.target]}):\n{cur}\n\n"
            f"Обсуждение правок:\n{convo}\n\nВыдай новый текст.")

    def _op():
        new = nodes._structural(st, None).complete(system, user)
        if req.target in ("chapter_plan", "chapter_dialogue"):
            chs = list(_state(thread_id).get("chapters") or [])
            if req.target == "chapter_plan":
                chs[idx].plan = new
            else:
                chs[idx].dialogue = new
            _patch(thread_id, run, {"chapters": chs})
            run.events.put({"type": "edited", "keys": ["chapters"]})
        elif req.target == "logline":
            lst = list(st.get("loglines") or [])
            sel = st.get("selected_logline")
            lst = [new if x == sel else x for x in lst] if sel in lst else [new, *lst]
            _patch(thread_id, run, {"loglines": lst, "selected_logline": new})
            run.events.put({"type": "edited", "keys": ["loglines"]})
        else:
            _patch(thread_id, run, {req.target: new})
            run.events.put({"type": "edited", "keys": [req.target]})

    return _bg_op(run, _op)


def _apply_project_target(thread_id: str, run, st: dict, target: str,
                          idx: int, new: str) -> None:
    """Записать новый контент в выбранный результат бота (для project_apply*)."""
    if target in ("chapter_plan", "chapter_dialogue"):
        chs = list(_state(thread_id).get("chapters") or [])
        if not (0 <= idx < len(chs)):
            return
        if target == "chapter_plan":
            chs[idx].plan = new
        else:
            chs[idx].dialogue = new
        _patch(thread_id, run, {"chapters": chs})
        run.events.put({"type": "edited", "keys": ["chapters"]})
    elif target == "logline":
        lst = list(st.get("loglines") or [])
        sel = st.get("selected_logline")
        lst = [new if x == sel else x for x in lst] if sel in lst else [new, *lst]
        _patch(thread_id, run, {"loglines": lst, "selected_logline": new})
        run.events.put({"type": "edited", "keys": ["loglines"]})
    else:
        _patch(thread_id, run, {target: new})
        run.events.put({"type": "edited", "keys": [target]})


@app.post("/api/runs/{thread_id}/project_apply_auto")
def project_apply_auto(thread_id: str, req: ProjectChatReq):
    """Ассистент САМ решает по итогу обсуждения, какой результат бота изменить,
    и вносит правку. Не управляет пайплайном/порядком — только содержимое."""
    from narrative import prompts  # noqa: PLC0415
    from narrative.state import ProjectAutoEdit  # noqa: PLC0415
    run = _get_run(thread_id)
    _busy(run)
    st = _state(thread_id)
    convo = nodes.to_english("\n".join(
        f"{'Нарративщик' if m.get('role') == 'user' else 'ИИ'}: "
        f"{m.get('content', '')}"
        for m in (req.messages or []) if m.get("role") in ("user", "assistant")
    ))
    system = (
        "Ты — ассистент-сценарист 18+. По итогу обсуждения реши, какой ОДИН "
        "результат бота нужно изменить (синопсис, персонажи, локации, логлайн, "
        "план или текст конкретной главы), и выдай его ПОЛНЫЙ новый вариант. "
        "Если по обсуждению менять нечего — changed=false, target=none. Меняй "
        "ТОЛЬКО содержимое, не трогай порядок глав и этапы. new_content — на "
        "английском, в том же формате, что был."
        + prompts.NAMES_RULE + prompts.CONTENT_POLICY
        + "\n\nКОНТЕКСТ ПРОЕКТА:\n" + _project_context(st)
    )
    user = f"Обсуждение:\n{convo}\n\nРеши, что изменить, и выдай новый текст."

    def _op():
        out = nodes._structural(st, None).structured(
            system, user, ProjectAutoEdit, temperature=0.3)
        if not out.changed or out.target == "none" or not out.new_content.strip():
            run.events.put({"type": "status", "status": run.status})
            return
        _apply_project_target(thread_id, run, st, out.target,
                              out.chapter_index, out.new_content)

    return _bg_op(run, _op)


@app.post("/api/runs/{thread_id}/chapter/{idx}/apply_chat")
def apply_chat(thread_id: str, idx: int, req: ChatReq):
    """Применить обсуждённое в чате к главе целиком (или не трогать).

    ИИ читает всю историю чата + главу и решает: внести согласованные правки
    или оставить как есть. Адалт-глава переписывается uncensored-моделью.
    """
    from narrative.state import ChatApplyOut  # noqa: PLC0415
    run = _get_run(thread_id)
    _busy(run)
    st = _state(thread_id)
    chapters = list(st.get("chapters") or [])
    if idx < 0 or idx >= len(chapters):
        raise HTTPException(400, "bad chapter index")
    ch = chapters[idx]
    convo = "\n".join(
        f"{'Нарративщик' if m.get('role') == 'user' else 'ИИ'}: {m.get('content','')}"
        for m in req.messages if m.get("role") in ("user", "assistant")
    )
    system = (
        "Ты — редактор-исполнитель визуальной новеллы 18+. На основе ОБСУЖДЕНИЯ "
        "с нарративщиком внеси в главу согласованные правки. Если из разговора "
        "НЕ следуют конкретные изменения — верни changed=false и не трогай текст. "
        "Сохрани формат ВН, теги статиков/анимаций, откровенность адалта.\n\n"
        f"Карточки:\n{st.get('characters','')}"
    )
    user = (
        f"ОБСУЖДЕНИЕ:\n{convo}\n\nТЕКУЩАЯ ГЛАВА {idx + 1} «{ch.title}»:\n"
        f"{ch.dialogue or ch.plan}\n\nПримени решения обсуждения к главе."
    )
    llm = nodes._adult() if ch.is_adult_point else nodes._structural(st, "dialogue")
    try:
        out = llm.structured(system, user, ChatApplyOut, temperature=0.6)
    except LLMLimitError as exc:
        raise HTTPException(429, _limit_detail(exc))
    except Exception as exc:
        raise HTTPException(502, f"не удалось применить (LLM): {str(exc)[:160]}")
    if out.changed and out.script.strip():
        ch.dialogue = out.script
        chapters[idx] = ch
        _patch(thread_id, run, {"chapters": chapters})
        run.events.put({"type": "edited", "keys": ["chapters"]})
    return {"changed": out.changed, "note": out.note}


# ---------- цикл ревизий главы: применить правки редактора + перепроверить ----

class RevisionReq(BaseModel):
    """Применение одной ревизии главы: чат по правкам учитывается при фиксе."""
    messages: list[dict] = []  # обсуждение этой ревизии (опционально)


def _open_findings(report) -> list:
    """Не отклонённые findings (их и исправляем)."""
    return [f for f in report.findings if f.status != "rejected"]


def _revision_worker(run: Run, idx: int, to_fix: list, convo: str):
    """ФОНОВО: точечная правка главы + перепроверка. LLM-правка адалта долгая
    (grok 30-150с), поэтому НЕ держим HTTP-запрос (иначе Next-прокси рвёт по
    таймауту → 500). Фронт поллит status: running → paused/error."""
    from narrative import live  # noqa: PLC0415
    thread_id = run.thread_id
    run.status = "running"
    run.error = ""
    run.events.put({"type": "status", "status": "running"})
    try:
        st = _state(thread_id)
        chapters = list(st.get("chapters") or [])
        if to_fix or convo.strip():
            # 1-2) ТОЧЕЧНАЯ правка (не перегенерация)
            convo = nodes.to_english(convo)  # RU-обсуждение → EN для агента
            st["chapter_idx"] = idx
            ch = nodes.patch_chapter(st, chapters[idx], idx, to_fix, extra=convo)
            if run.cancel:  # «Стоп» во время правки — правку не записываем
                raise _Cancelled()
            chapters[idx] = ch
            _patch(thread_id, run, {"chapters": chapters})
        if run.cancel:  # «Стоп» — перепроверку не запускаем
            raise _Cancelled()
        # 3) перепроверить → новый раунд редактора
        st2 = _state(thread_id)
        st2["chapter_idx"] = idx
        res = nodes.editor_node(st2)
        _patch(thread_id, run, res)
        run.status = "paused"
        run.events.put({"type": "edited",
                        "keys": ["chapters", "editor_reports"]})
        run.events.put({"type": "status", "status": "paused"})
    except _Cancelled:
        run.cancel = False
        run.status = "paused"
        run.events.put({"type": "status", "status": "paused"})
    except LLMLimitError as exc:  # #5: лимит → баннер выбора, прогресс цел
        run.limit = {"provider": exc.provider, "reset_at": exc.reset_at,
                     "kind": exc.kind, "message": str(exc)[:200]}
        run.status = "limit"
        run.events.put({"type": "status", "status": "limit",
                        "limit": run.limit})
    except Exception as exc:  # noqa: BLE001
        run.status = "error"
        run.error = f"Не удалось применить правки (LLM/провайдер): {str(exc)[:200]}"
        run.events.put({"type": "status", "status": "error", "error": run.error})
    finally:
        live.clear()
        run.gen = None


@app.post("/api/runs/{thread_id}/chapter/{idx}/apply_revision")
def apply_revision(thread_id: str, idx: int, req: RevisionReq):
    """Запускает один шаг цикла правок главы В ФОНЕ и сразу отвечает.

    Шаг: собрать НЕотклонённые замечания + обсуждение из чата → точечно
    исправить главу → редактор перепроверяет (новый раунд). Фронт поллит
    state до status=paused. Нарративщик жмёт, пока замечания не закроются.
    """
    run = _get_run(thread_id)
    _busy(run)
    st = _state(thread_id)
    chapters = list(st.get("chapters") or [])
    if idx < 0 or idx >= len(chapters):
        raise HTTPException(400, "bad chapter index")
    last = _last_report(st, idx)
    if last is None:
        raise HTTPException(400, "по этой главе ещё нет отчёта редактора")
    to_fix = _open_findings(last)
    convo_raw = "\n".join(
        f"{'Нарративщик' if m.get('role') == 'user' else 'ИИ'}: "
        f"{m.get('content', '')}"
        for m in (req.messages or []) if m.get("role") in ("user", "assistant")
    )
    # Нечего чинить (все замечания отклонены/закрыты) и обсуждения нет — НЕ
    # перезапускаем редактора вхолостую. Иначе свежий прогон editor_node может
    # придумать новые/повторные замечания и отправить главу на ревизию заново,
    # хотя нарративщик уже всё отклонил (баг: шаг 3 раньше вызывался безусловно).
    if not to_fix and not convo_raw.strip():
        return {"ok": True, "started": False,
                "note": "нечего чинить — все замечания уже закрыты"}
    # СИНХРОННО running до старта потока — иначе немедленный poll фронта ловит
    # старый «paused» и перестаёт поллить (та же гонка, что была в _bg_op).
    # Перевод RU→EN — внутри воркера: это LLM-вызов, ему не место в HTTP-запросе.
    with _LOCK:  # TOCTOU: двойной клик «Исправить всё» → два воркера
        if run.worker and run.worker.is_alive():
            raise HTTPException(409, "run is busy")
        run.cancel = False
        run.status = "running"
        run.error = ""
        run.worker = threading.Thread(
            target=_revision_worker, args=(run, idx, to_fix, convo_raw),
            daemon=True)
        run.worker.start()
    return {"ok": True, "started": True}


# ---------- откат на этап ----------

@app.post("/api/runs/{thread_id}/rollback")
def rollback(thread_id: str, req: RollbackReq):
    run = _get_run(thread_id)
    _busy(run)
    if req.stage not in ROLLBACK:
        raise HTTPException(400, f"rollback to {req.stage} не поддержан")
    as_node, clear = ROLLBACK[req.stage]
    patch: dict[str, Any] = {k: None for k in clear}
    if req.stage == "dialogue":
        # перезапуск написания: чистим тексты глав, структуру оставляем
        st = _state(thread_id)
        chs = list(st.get("chapters") or [])
        for c in chs:
            c.dialogue = c.adult_scene = c.translation = None
            c.statics, c.anims = [], []
            c.adult_statics, c.adult_anims = [], []
        # фаза заново «content»: сначала весь контент, редактор потом
        patch = {"chapters": chs, "chapter_idx": 0, "phase": "content"}
    elif req.stage in ("structure", "characters", "synopsis"):
        patch["chapters"] = []  # очистить главы (не-reducer)
    if req.stage in ("structure", "characters", "synopsis"):
        patch["structure_done"] = False
        patch["phase"] = None
    _patch(thread_id, run, patch, as_node=as_node)
    # откат снимает «done/error» — встаём на паузу на откатанном этапе, чтобы в
    # UI появилась кнопка «Продолжить», а не висел финальный экран «Проект готов».
    run.status = "paused"
    run.error = ""
    run.events.put({"type": "rollback", "stage": req.stage})
    run.events.put({"type": "status", "status": "paused"})
    return {"ok": True, "next": req.stage}


# ---------- SSE + health ----------

@app.get("/api/runs/{thread_id}/events")
async def events(thread_id: str):
    run = _get_run(thread_id)

    async def gen():
        snap = _graph(run).get_state(_config(thread_id))
        yield _sse({"type": "state", "state": _serialize(snap.values or {})})
        yield _sse({"type": "status", "status": run.status, "next": list(snap.next)})
        while True:
            try:
                yield _sse(run.events.get_nowait())
            except queue.Empty:
                await asyncio.sleep(0.2)
                yield ": ping\n\n"

    return StreamingResponse(gen(), media_type="text/event-stream")


def _sse(obj: dict) -> str:
    return f"data: {json.dumps(obj, ensure_ascii=False)}\n\n"


@app.get("/api/health")
def health():
    return {"ok": True}


# ---------- перевод (#3): кнопка «на русский» для нарративщика ----------

class TranslateReq(BaseModel):
    text: str
    to: str = "Russian"        # "Russian" | "English" | язык


# язык → ISO-код для Google Translate
_GT_CODE = {"russian": "ru", "english": "en", "ru": "ru", "en": "en"}


def _google_translate(text: str, to_code: str) -> str:
    """Быстрый перевод через Google Translate (общий модуль narrative.gtrans)."""
    from narrative.gtrans import google_translate  # noqa: PLC0415
    return google_translate(text, to_code)


@app.post("/api/translate")
def translate_ep(req: TranslateReq):
    """Перевод произвольного текста (вывод на английском → русский для чтения).

    Сначала пробуем быстрый Google Translate (без ИИ); если недоступен —
    фолбэк на LLM-перевод (сохраняет ВН-формат)."""
    if not (req.text or "").strip():
        return {"text": req.text, "via": "noop"}
    code = _GT_CODE.get(req.to.lower(), req.to.lower()[:2])
    # 1) быстрый Google (без ретраев — на заблокированном IP не виснем)
    try:
        return {"text": _google_translate(req.text, code), "via": "google"}
    except Exception:  # noqa: BLE001 — сеть/блокировка/парсинг → фолбэк на ИИ
        pass
    # 2) фолбэк на LLM
    lang = {"ru": "Russian", "en": "English"}.get(req.to.lower(), req.to)
    try:
        return {"text": nodes.translate(req.text, lang), "via": "llm"}
    except LLMLimitError as exc:
        raise HTTPException(429, _limit_detail(exc))
    except Exception:  # noqa: BLE001
        # 3) оба недоступны (напр. удалённый сервер без доступа к Google и без
        # LLM-креды) → отдаём ОРИГИНАЛ с 200, чтобы UI не падал с ошибкой/пустотой
        return {"text": req.text, "via": "none"}


# ---------- dev: оверрайд системных промптов по шагам (#4) ----------

def _prompt_defaults() -> dict:
    from narrative import prompts as _p  # noqa: PLC0415
    return {
        "logline": _p.LOGLINE, "synopsis": _p.SYNOPSIS,
        "characters": _p.CHARACTERS, "locations": _p.LOCATIONS,
        "structure": _p.STRUCTURE, "structure_editor": _p.STRUCTURE_EDITOR,
        "dialogue": _p.DIALOGUE, "adult": _p.ADULT,
        "editor": _p.EDITOR, "translation": _p.TRANSLATION,
    }


@app.get("/api/prompts")
def get_prompts():
    """Дефолтные системные промпты по шагам (для dev-редактора промптов)."""
    return {"defaults": _prompt_defaults()}


class PromptOverrideReq(BaseModel):
    stage: str
    text: str = ""        # пустой → сброс к дефолту


@app.post("/api/runs/{thread_id}/prompt_override")
def set_prompt_override(thread_id: str, req: PromptOverrideReq):
    """Задать/сбросить dev-оверрайд промпта шага для этого прогона."""
    run = _get_run(thread_id)
    st = _state(thread_id)
    ov = dict(st.get("prompt_overrides") or {})
    if req.text.strip():
        ov[req.stage] = req.text
    else:
        ov.pop(req.stage, None)
    _patch(thread_id, run, {"prompt_overrides": ov})
    return {"ok": True, "overrides": list(ov.keys())}


# ---------- экспорт готового проекта (скачиваемый текст) ----------

@app.get("/api/runs/{thread_id}/export")
def export_project(thread_id: str, fmt: str = "txt"):
    """Собирает готовую новеллу в один документ для скачивания.

    fmt=txt — чистый текст; fmt=md — с заголовками. Включает план (кратко),
    основной текст главы (диалоги+адалт) и перевод, если он есть.
    """
    st = _state(thread_id)
    chapters = list(st.get("chapters") or [])
    md = fmt == "md"

    from narrative import gtrans  # noqa: PLC0415
    # Группировка ПО ЯЗЫКАМ: сначала ВСЕ главы на одном языке, потом на следующем.
    parts: list[str] = []

    def _lang_block(lang_label: str, get_text) -> None:
        bodies = [(ch, (get_text(ch) or "").strip()) for ch in chapters]
        if not any(b for _, b in bodies):
            return
        parts.append(f"\n# {lang_label}" if md else f"\n\n########## {lang_label} ##########")
        for ch, body in bodies:
            if not body:
                continue
            head = f"Глава {ch.index + 1}. {ch.title}"
            parts.append(f"\n## {head}" if md else f"\n\n=== {head} ===")
            parts.append(body)

    # оригинал (English), затем каждый целевой язык
    _lang_block("English (оригинал)", lambda ch: ch.dialogue)
    for code in gtrans.TARGET_CODES:
        _lang_block(gtrans.NAME_BY_CODE.get(code, code),
                    lambda ch, c=code: (getattr(ch, "translations", None) or {}).get(c))

    text = "\n".join(parts).strip() + "\n"
    safe = "novella"
    fname = f"{safe}.{ 'md' if md else 'txt'}"
    return {"filename": fname, "text": text, "chapters": len(chapters)}


@app.get("/api/runs/{thread_id}/export.docx")
def export_project_docx(thread_id: str):
    """Собрать новеллу в .docx (бинарь) — заголовки глав + текст + перевод."""
    import io  # noqa: PLC0415
    try:
        from docx import Document  # noqa: PLC0415
    except ImportError:
        # деплой без зависимости → 501 с понятным текстом, а не голый 500
        raise HTTPException(
            501, "python-docx не установлен на сервере — обнови зависимости "
                 "(pip install -r requirements.txt / пересобери Docker)")
    from docx.oxml import OxmlElement  # noqa: PLC0415
    from docx.oxml.ns import qn  # noqa: PLC0415
    st = _state(thread_id)
    chapters = list(st.get("chapters") or [])

    from narrative import gtrans  # noqa: PLC0415
    # Группировка ПО ЯЗЫКАМ: заголовок языка (level 1), под ним все главы (level 2).
    doc = Document()

    def _add_body(body: str) -> None:
        """Текст главы ОДНИМ параграфом: строки через <w:br/> внутри одного run.
        При параграфе-на-строку (8 глав × 27 языков × ~500 строк ≈ 100k+ вызовов
        add_paragraph) сборка висла минутами и валила прокси по таймауту (500)."""
        p = doc.add_paragraph()
        r = p.add_run()._r
        for j, ln in enumerate(body.split("\n")):
            if j:
                r.append(OxmlElement("w:br"))
            t = OxmlElement("w:t")
            t.text = ln
            t.set(qn("xml:space"), "preserve")
            r.append(t)

    def _lang_section(lang_label: str, get_text) -> None:
        bodies = [(ch, (get_text(ch) or "").strip()) for ch in chapters]
        if not any(b for _, b in bodies):
            return
        doc.add_heading(lang_label, level=1)
        for ch, body in bodies:
            if not body:
                continue
            doc.add_heading(f"Глава {ch.index + 1}. {ch.title}", level=2)
            _add_body(body)

    _lang_section("English (оригинал)", lambda ch: ch.dialogue)
    for code in gtrans.TARGET_CODES:
        _lang_section(gtrans.NAME_BY_CODE.get(code, code),
                      lambda ch, c=code: (getattr(ch, "translations", None) or {}).get(c))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    headers = {"Content-Disposition": 'attachment; filename="novella.docx"'}
    return StreamingResponse(
        buf, headers=headers,
        media_type="application/vnd.openxmlformats-officedocument."
                   "wordprocessingml.document")


# ---------- Claude по подписке (#4 расширение): статус / тумблер / токен ----------

def _claude_login_status() -> dict:
    """Авторизован ли Claude (подписка). Источник: env-токен или credentials.json."""
    import json
    import os as _os
    import time
    if _os.environ.get("ANTHROPIC_API_KEY"):
        return {"authorized": True, "via": "api_key",
                "warn": "ANTHROPIC_API_KEY перебивает подписку — это платный API"}
    if _os.environ.get("CLAUDE_CODE_OAUTH_TOKEN"):
        return {"authorized": True, "via": "oauth_token"}
    path = _os.path.expanduser("~/.claude/.credentials.json")
    try:
        with open(path) as fh:
            o = json.load(fh)
        o = o.get("claudeAiOauth") or o
        exp = o.get("expiresAt")
        valid = True
        when = None
        if exp:
            secs = exp / 1000 if exp > 1e12 else exp
            valid = secs > time.time()
            when = time.strftime("%Y-%m-%d %H:%M", time.localtime(secs))
        return {"authorized": valid, "via": "credentials",
                "expires": when, "sub": o.get("subscriptionType")}
    except Exception:
        return {"authorized": False, "via": None}


@app.get("/api/claude/status")
def claude_status():
    import os as _os
    st = _claude_login_status()
    st["enabled"] = _os.environ.get("USE_CLAUDE_SUBSCRIPTION") == "1"
    st["model"] = _os.environ.get("CLAUDE_SUB_MODEL", "sonnet")
    return st


@app.get("/api/claude/usage")
def claude_usage():
    """Проактивный остаток лимита подписки Claude — из последнего RateLimitEvent
    SDK (utilization 0..1, время сброса, статус). None-поля = SDK ещё не сообщил."""
    from narrative.claude_sub import last_rate_limit  # noqa: PLC0415
    rl = last_rate_limit()
    return {"rate_limit": rl}


class SubReq(BaseModel):
    enabled: bool
    model: Optional[str] = None


@app.post("/api/claude/subscription")
def claude_subscription(req: SubReq):
    """Тумблер: все НЕ-адалт боты через подписку Claude + выбор модели."""
    import os as _os
    _os.environ["USE_CLAUDE_SUBSCRIPTION"] = "1" if req.enabled else "0"
    if req.model:
        _os.environ["CLAUDE_SUB_MODEL"] = req.model
    return claude_status()


class TokenReq(BaseModel):
    token: str


@app.post("/api/claude/token")
def claude_token(req: TokenReq):
    """Вставить долгоживущий токен из `claude setup-token` (необязательно —
    если есть валидный логин Claude Code, токен не нужен). Пишем в процесс."""
    import os as _os
    tok = req.token.strip()
    if not tok:
        raise HTTPException(400, "пустой токен")
    _os.environ["CLAUDE_CODE_OAUTH_TOKEN"] = tok
    return claude_status()


# ---------- OAuth-подключение подписки прямо из веба ----------
# Тот же флоу, что `claude setup-token`: открыть authorize-страницу в новой
# вкладке → пользователь логинится → копирует выданный code → мы меняем его на
# токен (PKCE). Публичный client_id Claude Code.

_CLAUDE_CLIENT_ID = "9d1c250a-e61b-44d9-88ed-5944d1962f5e"
_CLAUDE_REDIRECT = "https://console.anthropic.com/oauth/code/callback"
_CLAUDE_SCOPE = "org:create_api_key user:profile user:inference"
_PKCE: dict[str, str] = {}  # verifier последнего запроса (один за раз)


@app.get("/api/claude/auth_url")
def claude_auth_url():
    """Сгенерировать ссылку авторизации (PKCE) — для подключения или смены
    аккаунта. Возвращаем всегда (даже если уже залогинен — чтобы можно было
    переподключить другой аккаунт)."""
    import base64
    import hashlib
    import secrets
    from urllib.parse import urlencode
    verifier = secrets.token_urlsafe(64)
    challenge = base64.urlsafe_b64encode(
        hashlib.sha256(verifier.encode()).digest()).rstrip(b"=").decode()
    _PKCE["v"] = verifier
    params = {
        "code": "true", "client_id": _CLAUDE_CLIENT_ID,
        "response_type": "code", "redirect_uri": _CLAUDE_REDIRECT,
        "scope": _CLAUDE_SCOPE, "code_challenge": challenge,
        "code_challenge_method": "S256", "state": verifier,
    }
    return {"authorized": False,
            "url": "https://claude.ai/oauth/authorize?" + urlencode(params)}


class CodeReq(BaseModel):
    code: str


@app.post("/api/claude/exchange")
def claude_exchange(req: CodeReq):
    """Обменять выданный code (формат code#state) на токен подписки."""
    import httpx
    raw = req.code.strip()
    if not raw:
        raise HTTPException(400, "пустой код")
    code, _, state = raw.partition("#")
    verifier = _PKCE.get("v")
    if not verifier:
        raise HTTPException(400, "сначала открой ссылку авторизации")
    try:
        r = httpx.post(
            "https://console.anthropic.com/v1/oauth/token",
            json={
                "grant_type": "authorization_code", "code": code,
                "state": state or verifier, "client_id": _CLAUDE_CLIENT_ID,
                "redirect_uri": _CLAUDE_REDIRECT, "code_verifier": verifier,
            }, timeout=30)
    except Exception as exc:
        raise HTTPException(502, f"сеть/обмен не удался: {str(exc)[:140]}")
    if r.status_code != 200:
        raise HTTPException(502, f"обмен не удался: {r.status_code} "
                            f"{r.text[:140]}")
    tok = (r.json() or {}).get("access_token")
    if not tok:
        raise HTTPException(502, "в ответе нет access_token")
    import os as _os
    _os.environ["CLAUDE_CODE_OAUTH_TOKEN"] = tok
    _PKCE.pop("v", None)
    return claude_status()
